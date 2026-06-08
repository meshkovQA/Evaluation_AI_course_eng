import os
import io
from typing import Dict, List, Optional
from pathlib import Path
import mimetypes

# PDF parsing
import PyPDF2
from pdfplumber import PDF

# Document parsing
from docx import Document
import openpyxl

# Text and other formats
import csv
import json
from bs4 import BeautifulSoup
import markdown


class FileParser:
    """Class for parsing various file formats"""

    SUPPORTED_FORMATS = {
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'application/vnd.ms-word': 'doc',
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx',
        'application/vnd.ms-excel': 'xls',
        'text/plain': 'txt',
        'text/csv': 'csv',
        'application/json': 'json',
        'text/html': 'html',
        'text/markdown': 'md',
        'text/x-markdown': 'md'
    }

    EXTENSION_TO_TYPE = {
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'doc',
        '.xlsx': 'xlsx',
        '.xls': 'xls',
        '.txt': 'txt',
        '.csv': 'csv',
        '.json': 'json',
        '.html': 'html',
        '.htm': 'html',
        '.md': 'md',
        '.markdown': 'md'
    }

    def __init__(self):
        mimetypes.init()

    def get_file_type(self, file_path: str) -> Optional[str]:
        """Determines file type by MIME type or extension"""
        mime_type, _ = mimetypes.guess_type(file_path)
        if mime_type and mime_type in self.SUPPORTED_FORMATS:
            return self.SUPPORTED_FORMATS[mime_type]

        extension = Path(file_path).suffix.lower()
        return self.EXTENSION_TO_TYPE.get(extension)

    def is_supported(self, file_path: str) -> bool:
        """Checks if the file format is supported"""
        file_type = self.get_file_type(file_path)
        return file_type is not None

    def parse_file(self, file_path: str) -> Dict[str, any]:
        """
        Parses a file and returns the extracted text and metadata

        Returns:
            Dict with keys:
            - text: extracted text
            - metadata: file metadata
            - file_type: file type
            - success: parsing status
            - error: error message (if any)
        """
        result = {
            'text': '',
            'metadata': {},
            'file_type': None,
            'success': False,
            'error': None
        }

        try:
            if not os.path.exists(file_path):
                result['error'] = f"File not found: {file_path}"
                return result

            file_type = self.get_file_type(file_path)
            if not file_type:
                result['error'] = f"Unsupported file format: {file_path}"
                return result

            result['file_type'] = file_type
            result['metadata'] = self._get_basic_metadata(file_path)

            if file_type == 'pdf':
                result['text'] = self._parse_pdf(file_path)
            elif file_type == 'docx':
                result['text'] = self._parse_docx(file_path)
            elif file_type in ['txt', 'md']:
                result['text'] = self._parse_text_file(file_path)
            elif file_type == 'csv':
                result['text'] = self._parse_csv(file_path)
            elif file_type == 'json':
                result['text'] = self._parse_json(file_path)
            elif file_type == 'html':
                result['text'] = self._parse_html(file_path)
            elif file_type in ['xlsx', 'xls']:
                result['text'] = self._parse_excel(file_path)
            else:
                result['error'] = f"Parser for type {file_type} is not implemented"
                return result

            result['success'] = True

        except Exception as e:
            result['error'] = f"Error parsing file: {str(e)}"

        return result

    def _get_basic_metadata(self, file_path: str) -> Dict[str, any]:
        """Gets basic file metadata"""
        file_stat = os.stat(file_path)
        return {
            'filename': os.path.basename(file_path),
            'file_size': file_stat.st_size,
            'created_at': file_stat.st_ctime,
            'modified_at': file_stat.st_mtime,
            'file_extension': Path(file_path).suffix
        }

    def _parse_pdf(self, file_path: str) -> str:
        """Parses a PDF file"""
        text = ""

        # Try pdfplumber first (better handles tables)
        try:
            with PDF.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
        except:
            # Fall back to PyPDF2 if pdfplumber fails
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n\n"

        return text.strip()

    def _parse_docx(self, file_path: str) -> str:
        """Parses a DOCX file"""
        doc = Document(file_path)
        text = ""

        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"

        return text.strip()

    def _parse_text_file(self, file_path: str) -> str:
        """Parses text files (txt, md)"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            content = file.read()

        # If markdown, convert to plain text
        if file_path.endswith('.md'):
            content = markdown.markdown(content)
            soup = BeautifulSoup(content, 'html.parser')
            content = soup.get_text()

        return content.strip()

    def _parse_csv(self, file_path: str) -> str:
        """Parses a CSV file"""
        text = ""

        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            csv_reader = csv.reader(file)
            for row in csv_reader:
                text += " | ".join(row) + "\n"

        return text.strip()

    def _parse_json(self, file_path: str) -> str:
        """Parses a JSON file"""
        with open(file_path, 'r', encoding='utf-8') as file:
            data = json.load(file)

        return json.dumps(data, ensure_ascii=False, indent=2)

    def _parse_html(self, file_path: str) -> str:
        """Parses an HTML file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            content = file.read()

        soup = BeautifulSoup(content, 'html.parser')

        # Remove scripts and styles
        for script in soup(["script", "style"]):
            script.decompose()

        return soup.get_text().strip()

    def _parse_excel(self, file_path: str) -> str:
        """Parses an Excel file"""
        workbook = openpyxl.load_workbook(file_path, read_only=True)
        text = ""

        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text += f"Sheet: {sheet_name}\n"

            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(
                    [str(cell) if cell is not None else "" for cell in row])
                if row_text.strip():
                    text += row_text + "\n"

            text += "\n"

        workbook.close()
        return text.strip()


def parse_document(file_path: str) -> Dict[str, any]:
    """
    Convenience function for parsing a document

    Args:
        file_path: path to the file

    Returns:
        Dict with parsing result
    """
    parser = FileParser()
    return parser.parse_file(file_path)
